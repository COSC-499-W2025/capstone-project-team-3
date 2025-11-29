"""pytest tests for document parser: text, markdown, docx and pdf."""
import json
from pathlib import Path
from app.utils.non_code_parsing.document_parser import parse_documents_to_json, parsed_input_text
from app.utils.non_code_parsing.file_size_checker import is_file_too_large, MAX_FILE_SIZE_MB


def test_parse_text_and_markdown(tmp_path=None):
    # provide a tmp_path when running directly
    if tmp_path is None:
        from tempfile import mkdtemp
        tmp_path = Path(mkdtemp())

    # text file
    txt = tmp_path / "a.txt"
    txt.write_text("Hello txt")

    # markdown file
    md = tmp_path / "b.md"
    md.write_text("# Title\ncontent")

    output = tmp_path / "out.json"
    parsed = parse_documents_to_json([str(txt), str(md)], str(output))

    files = {f["name"]: f for f in parsed["files"]}
    assert files["a.txt"]["success"] is True
    assert files["a.txt"]["content"] == "Hello txt"
    assert files["b.md"]["success"] is True
    assert "Title" in files["b.md"]["content"] or "content" in files["b.md"]["content"]


def test_parse_docx(tmp_path=None):
    if tmp_path is None:
        from tempfile import mkdtemp
        tmp_path = Path(mkdtemp())

    try:
        from docx import Document
    except Exception:
        print('skipping docx test: python-docx not installed')
        return

    docx_path = tmp_path / "d.docx"
    doc = Document()
    doc.add_paragraph("Paragraph 1")
    doc.add_paragraph("Paragraph 2")
    doc.save(str(docx_path))

    output = tmp_path / "out.json"
    parsed = parse_documents_to_json([str(docx_path)], str(output))
    entry = parsed["files"][0]
    assert entry["success"] is True
    assert "Paragraph 1" in entry["content"]


def test_parse_pdf(tmp_path=None):
    if tmp_path is None:
        from tempfile import mkdtemp
        tmp_path = Path(mkdtemp())

    # create a PDF using reportlab (already in requirements)
    try:
        from reportlab.pdfgen import canvas
    except Exception:
        print('skipping pdf test: reportlab not installed')
        return

    pdf_path = tmp_path / "p.pdf"
    c = canvas.Canvas(str(pdf_path))
    c.drawString(100, 750, "PDF line 1")
    c.drawString(100, 730, "PDF line 2")
    c.save()

    output = tmp_path / "out.json"
    parsed = parse_documents_to_json([str(pdf_path)], str(output))
    entry = parsed["files"][0]
    assert entry["success"] is True
    assert "PDF line 1" in entry["content"] or "PDF line" in entry["content"]


def test_missing_file_and_unsupported_type(tmp_path=None):
    if tmp_path is None:
        from tempfile import mkdtemp
        tmp_path = Path(mkdtemp())

    # Missing file
    missing = tmp_path / "nope.txt"
    output = tmp_path / "out.json"
    returned = parse_documents_to_json([str(missing)], str(output))
    assert len(returned["files"]) == 1
    entry = returned["files"][0]
    assert entry["success"] is False
    assert "does not exist" in entry["error"].lower()

    # Unsupported type
    test_file = tmp_path / "data.bin"
    test_file.write_bytes(b"\x00\x01\x02")
    returned2 = parse_documents_to_json([str(test_file)], str(output))
    entry2 = returned2["files"][0]
    assert entry2["success"] is False
    assert "unsupported" in entry2["error"].lower()

def test_file_size_limit(tmp_path=None):
    if tmp_path is None:
        from tempfile import mkdtemp
        tmp_path = Path(mkdtemp())
    # Create a file slightly larger than MAX_FILE_SIZE_MB
    too_big_path = tmp_path / "big.txt"
    # write enough bytes to exceed MAX_FILE_SIZE_MB
    size_bytes = (MAX_FILE_SIZE_MB + 1) * 1024 * 1024  # MB -> bytes
    with open(too_big_path, "wb") as f:
        f.write(b"0" * size_bytes)

    output = tmp_path / "out.json"
    parsed = parse_documents_to_json([str(too_big_path)], str(output))

    entry = parsed["files"][0]
    assert entry["success"] is False
    assert "exceeds" in entry["error"].lower()



def test_parsed_input_text_basic(tmp_path=None):
    """Test parsed_input_text function with text and markdown files."""
    if tmp_path is None:
        from tempfile import mkdtemp
        tmp_path = Path(mkdtemp())
    
    # Create test files
    txt_file = tmp_path / "test.txt"
    txt_file.write_text("Hello world")
    
    md_file = tmp_path / "readme.md"
    md_file.write_text("# Title\nContent here")
    
    # Test parsing
    result = parsed_input_text([str(txt_file), str(md_file)])
    
    # Verify structure
    assert "parsed_files" in result
    assert len(result["parsed_files"]) == 2
    
    # Verify all required fields exist
    for file_result in result["parsed_files"]:
        assert "path" in file_result
        assert "name" in file_result
        assert "type" in file_result
        assert "content" in file_result
        assert "success" in file_result
        assert "error" in file_result


def test_parsed_input_text_missing_file(tmp_path=None):
    """Test parsed_input_text with missing files."""
    result = parsed_input_text(["/nonexistent/file.txt"])
    
    assert len(result["parsed_files"]) == 1
    file_result = result["parsed_files"][0]
    assert file_result["success"] is False
    assert "does not exist" in file_result["error"].lower()


def test_parsed_input_text_empty_list(tmp_path=None):
    """Test parsed_input_text with empty file list."""
    result = parsed_input_text([])
    assert result == {"parsed_files": []}


if __name__ == "__main__":
    # run tests directly
    test_parse_text_and_markdown()
    print('test_parse_text_and_markdown passed')
    try:
        test_parse_docx()
        print('test_parse_docx passed')
    except Exception as e:
        print('test_parse_docx failed:', e)
    try:
        test_parse_pdf()
        print('test_parse_pdf passed')
    except Exception as e:
        print('test_parse_pdf failed:', e)
    
    # Test new parsed_input_text function
    test_parsed_input_text_basic()
    print('test_parsed_input_text_basic passed')
    
    test_parsed_input_text_missing_file()
    print('test_parsed_input_text_missing_file passed')
    
    test_parsed_input_text_empty_list()
    print('test_parsed_input_text_empty_list passed')
